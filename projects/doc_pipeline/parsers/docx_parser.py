"""
DOCX Parser — Extracts text from Word documents using python-docx
===================================================================
"""

import io
from docx import Document


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        file_bytes: Raw bytes of the DOCX file

    Returns:
        Extracted text as a single string

    Raises:
        ValueError: If the DOCX cannot be read or is empty
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Failed to open DOCX: {str(e)}")

    paragraphs = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level_num = int(level)
                    prefix = "#" * level_num
                    paragraphs.append(f"{prefix} {text}")
                except ValueError:
                    paragraphs.append(text)
            else:
                paragraphs.append(text)

    # Extract tables
    for i, table in enumerate(doc.tables):
        table_text = [f"\n--- Table {i + 1} ---"]
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(cells))
        paragraphs.append("\n".join(table_text))

    if not paragraphs:
        raise ValueError("DOCX contains no extractable text")

    return "\n\n".join(paragraphs)
